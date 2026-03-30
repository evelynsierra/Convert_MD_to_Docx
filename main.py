import argparse
import subprocess
import shutil
import sys
import re
from pathlib import Path

DEFAULT_INPUT_FOLDER  = "docs"
DEFAULT_OUTPUT_FOLDER = "docx_output"


# ──────────────────────────────────────────────
# Method 1: pandoc (recommended, best fidelity)
# ──────────────────────────────────────────────

def convert_with_pandoc(md_path: Path, docx_path: Path) -> bool:
    if not shutil.which("pandoc"):
        print("  [!] pandoc not found. Install from https://pandoc.org/installing.html")
        return False

    cmd = ["pandoc", str(md_path), "-o", str(docx_path),
           "--from", "markdown", "--to", "docx", "--standalone"]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"  [!] pandoc error: {result.stderr.strip()}")
        return False
    return True


# ──────────────────────────────────────────────
# Method 2: Pure Python (python-docx fallback)
# ──────────────────────────────────────────────

def convert_with_python_docx(md_path: Path, docx_path: Path) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        print("  [!] python-docx not installed. Run: pip install python-docx")
        return False

    doc = Document()

    def set_code_style(para):
        para.style = "Normal"
        for run in para.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def add_inline_runs(para, text):
        pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
        parts = re.split(pattern, text)
        for chunk in parts:
            if not chunk:
                continue
            if chunk.startswith("**") and chunk.endswith("**"):
                run = para.add_run(chunk[2:-2])
                run.bold = True
            elif chunk.startswith("*") and chunk.endswith("*"):
                run = para.add_run(chunk[1:-1])
                run.italic = True
            elif chunk.startswith("`") and chunk.endswith("`"):
                run = para.add_run(chunk[1:-1])
                run.font.name = "Courier New"
                run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            else:
                if not re.match(r'\*\*.*\*\*|\*.*\*|`.*`', chunk):
                    para.add_run(chunk)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ""

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lang = line.strip()[3:].strip()
                code_lines = []
            else:
                in_code_block = False
                if code_lang:
                    set_code_style(doc.add_paragraph(f"[{code_lang}]"))
                for cl in code_lines:
                    set_code_style(doc.add_paragraph(cl if cl else " "))
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if re.match(r'^(\-{3,}|\*{3,}|_{3,})\s*$', line):
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            doc.add_heading(m.group(2).strip(), level=min(len(m.group(1)), 9))
            i += 1
            continue

        if line.startswith("> "):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Pt(36)
            run = p.add_run(line[2:].strip())
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        m = re.match(r'^\s*[\-\*\+]\s+(.*)', line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, m.group(1).strip())
            i += 1
            continue

        m = re.match(r'^\s*\d+\.\s+(.*)', line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, m.group(1).strip())
            i += 1
            continue

        if line.strip() == "":
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, line.strip())
        i += 1

    doc.save(str(docx_path))
    return True


# ──────────────────────────────────────────────
# Convert dispatcher
# ──────────────────────────────────────────────

def convert(md_path: Path, docx_path: Path, method: str) -> None:
    if not md_path.exists():
        print(f"  [!] File not found: {md_path}")
        return

    docx_path.parent.mkdir(parents=True, exist_ok=True)

    print(f"  Converting : {md_path}")
    print(f"  Output     : {docx_path}")

    if method == "pandoc":
        ok = convert_with_pandoc(md_path, docx_path)
    elif method == "python":
        ok = convert_with_python_docx(md_path, docx_path)
    else:
        ok = convert_with_pandoc(md_path, docx_path)
        if not ok:
            print("  Falling back to python-docx …")
            ok = convert_with_python_docx(md_path, docx_path)

    print(f"  {'✓ Saved' if ok else '✗ Failed'}\n")


# ──────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown (.md) files to Word (.docx)"
    )
    parser.add_argument(
        "inputs", nargs="*",
        help="One or more .md file paths. If omitted, all .md in --folder are converted."
    )
    parser.add_argument(
        "--folder", default=DEFAULT_INPUT_FOLDER,
        help=f'Input folder to scan (default: "{DEFAULT_INPUT_FOLDER}")'
    )
    parser.add_argument(
        "--out-folder", default=DEFAULT_OUTPUT_FOLDER,
        help=f'Output folder for .docx files (default: "{DEFAULT_OUTPUT_FOLDER}"). '
             'Created automatically if it does not exist.'
    )
    parser.add_argument(
        "-o", "--output",
        help="Exact output .docx path (only valid for single-file conversion)"
    )
    parser.add_argument(
        "--method", choices=["auto", "pandoc", "python"], default="auto",
        help="'auto' (default): try pandoc, fall back to python-docx. "
             "'pandoc': require pandoc. 'python': use python-docx only."
    )
    args = parser.parse_args()

    if args.inputs:
        inputs = [Path(p) for p in args.inputs]
    else:
        folder = Path(args.folder)
        if not folder.exists():
            print(f'[!] Input folder not found: "{folder}"')
            print(f'    Make sure "{DEFAULT_INPUT_FOLDER}/" exists next to this script,')
            print(f'    or pass a custom path: --folder "path/to/folder"')
            sys.exit(1)
        inputs = sorted(folder.glob("*.md"))
        if not inputs:
            print(f'[!] No .md files found in "{folder}"')
            sys.exit(1)
        print(f'\nFound {len(inputs)} .md file(s) in "{folder}":\n')
        for f in inputs:
            print(f"  • {f.name}")
        print()

    if args.output and len(inputs) > 1:
        print("[!] -o/--output can only be used when converting a single file.")
        sys.exit(1)

    out_folder = Path(args.out_folder)

    for md_path in inputs:
        if args.output and len(inputs) == 1:
            docx_path = Path(args.output)
        else:
            docx_path = out_folder / md_path.with_suffix(".docx").name
        convert(md_path, docx_path, args.method)

    print("Done.")


if __name__ == "__main__":
    main()